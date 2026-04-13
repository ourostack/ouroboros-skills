#!/usr/bin/env python3
"""Post-process a pandoc-generated .docx to improve styling.

Adds:
- Table grid borders with shaded header rows
- Code block grey backgrounds with Consolas font
- Configurable page margins

Usage:
  python3 style_docx.py output.docx
  python3 style_docx.py output.docx --margins 0.5
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def style_tables(doc: Document) -> None:
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "BFBFBF")
            borders.append(el)

        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(borders)

        # Header row: blue shading + bold
        for cell in table.rows[0].cells:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9E2F3")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        # Compact cell spacing
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(9)


def style_code_blocks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        is_code = "Source" in style_name or "Verbatim" in style_name

        if not is_code and paragraph.runs:
            fonts = [r.font.name for r in paragraph.runs if r.font.name]
            is_code = (
                all(f in ("Courier New", "Courier", "Consolas", "Menlo") for f in fonts)
                if fonts
                else False
            )

        if is_code:
            # Grey background
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F2F2")
            shading.set(qn("w:val"), "clear")
            paragraph._p.get_or_add_pPr().append(shading)

            # Subtle border
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "4")
                el.set(qn("w:color"), "DCDCDC")
                pBdr.append(el)
            existing = pPr.find(qn("w:pBdr"))
            if existing is not None:
                pPr.remove(existing)
            pPr.append(pBdr)

            # Font + spacing
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.left_indent = Inches(0.15)


def set_margins(doc: Document, margin_inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", help="Path to the .docx file to style")
    parser.add_argument(
        "--margins", type=float, default=0.5, help="Page margins in inches (default: 0.5)"
    )
    args = parser.parse_args()

    path = Path(args.docx).expanduser().resolve()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")

    doc = Document(str(path))
    style_tables(doc)
    style_code_blocks(doc)
    set_margins(doc, args.margins)
    doc.save(str(path))
    print(path)


if __name__ == "__main__":
    main()
