#!/usr/bin/env python3
"""Convert a markdown file into a simple .docx document.

Supported markdown features:
- # / ## headings
- blank lines as paragraph breaks
- inline links [text](url)
- bold **text**
- inline code `code`

Usage:
  python3 md_to_docx.py /absolute/path/to/source.md
  python3 md_to_docx.py /absolute/path/to/source.md --output /absolute/path/to/output.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    new_run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_inline(paragraph, line: str) -> None:
    i = 0
    while i < len(line):
        earliest = None
        chosen = None
        for kind, regex in (("link", LINK_RE), ("bold", BOLD_RE), ("code", CODE_RE)):
            match = regex.search(line, i)
            if match and (earliest is None or match.start() < earliest.start()):
                earliest = match
                chosen = kind

        if earliest is None:
            paragraph.add_run(line[i:])
            break

        if earliest.start() > i:
            paragraph.add_run(line[i : earliest.start()])

        if chosen == "link":
            add_hyperlink(paragraph, earliest.group(1), earliest.group(2))
        elif chosen == "bold":
            run = paragraph.add_run(earliest.group(1))
            run.bold = True
        elif chosen == "code":
            run = paragraph.add_run(earliest.group(1))
            run.font.name = "Menlo"

        i = earliest.end()


def convert_markdown_to_docx(source: Path, output: Path) -> None:
    text = source.read_text()
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    for line in text.splitlines():
        if not line.strip():
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue

        paragraph = document.add_paragraph()
        render_inline(paragraph, line)

    document.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", help="Path to the source markdown file")
    parser.add_argument(
        "--output",
        help="Optional output .docx path. Defaults to the source path with a .docx suffix.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source = Path(args.source).expanduser().resolve()
    if not source.exists():
        raise SystemExit(f"Source file not found: {source}")
    if source.suffix.lower() != ".md":
        raise SystemExit(f"Source file must be markdown (.md): {source}")

    output = (
        Path(args.output).expanduser().resolve()
        if args.output
        else source.with_suffix(".docx")
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    convert_markdown_to_docx(source, output)
    print(output)


if __name__ == "__main__":
    main()
